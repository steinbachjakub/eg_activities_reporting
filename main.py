# TODO: Duplicate activities (similar / same name and same date)
# TODO: Restore metrics from the previous reports
# TODO: Restore the report generator
# TODO: Country reports
# TODO: Example of section reports for universities
# TODO: Connection with Google Looker Studio

import requests
import pandas as pd
from pathlib import Path
from matplotlib import pyplot as plt
from datetime import datetime, date
from dateutil.relativedelta import relativedelta
from bs4 import BeautifulSoup as BSoup
import geopandas as gpd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from matplotlib.lines import Line2D

#######################################################################################################################
######################################### !!!NEED TO REFACTOR!!! ######################################################
#######################################################################################################################
#
#
# def generate_doc_intro(document, styles, date_from, date_to, event_count, part_tot, part_mean):
#
#     title_style, header_style, text_style = styles
#     document.add_paragraph('Monthly Activities Report', style=title_style)
#     document.add_paragraph(
#         f'Dear network, this is the regular report of events submitted in ESN Activities. Below, you can find some highlights (for the period of {date_from} to {date_to}).',
#         style=text_style)
#     p = document.add_paragraph(style=text_style)
#     p.add_run(
#         "Are you curious about how your National Organisation or your local sections fare? Would you like us to create a similar report for you? No problem! Just contact us at ").bold = True
#     p.add_run("social-impact@esn.org").italic = True
#     p.add_run(".").bold = True
#     document.add_paragraph('Few Numbers First', style=header_style)
#     document.add_paragraph(
#         f"During the period of {date_from} to {date_to}, the total of {event_count:,} events were organised. All of our events combined, we reached {part_tot:,} participants joining our events which results in the average of {part_mean:.2f} participants per one event.",
#         style=text_style)
#     document.add_paragraph(
#         "You can see the comparison of event and participant numbers in the maps below.",
#         style=text_style)
#
#
# def change_styles(document_styles):
#     title_style = document_styles.add_style("TitleStyle", WD_STYLE_TYPE.PARAGRAPH)
#     title_style.font.name = "Kelson Sans"
#     title_style.font.size = Pt(24)
#     title_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
#     header_style = document_styles.add_style("HeaderStyle", WD_STYLE_TYPE.PARAGRAPH)
#     header_style.font.name = "Lato"
#     header_style.font.size = Pt(16)
#     header_style.font.bold = True
#     header_style.font.color.rgb = RGBColor(0x2E, 0x31, 0x92)
#     header_style.paragraph_format.keep_with_next = True
#     text_style = document_styles.add_style("TextStyle", WD_STYLE_TYPE.PARAGRAPH)
#     text_style.font.name = "Lato"
#     text_style.font.size = Pt(11)
#     text_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
#     text_style.paragraph_format.keep_with_next = True
#     text_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
#
#     return title_style, header_style, text_style
#
#
# def generate_doc_section(document, styles, text=None, title=None):
#     title_style, header_style, text_style = styles
#     if (text is not None) and (title is not None):
#         document.add_paragraph(title, style=header_style)
#         document.add_paragraph(text, style=text_style)
#     document.add_picture("fig_temp.png", height=Cm(8))
#     document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
#     document.paragraphs[-1].paragraph_format.keep_with_next = False
#
#
# def generate_doc_outro(document, styles):
#     title_style, header_style, text_style = styles
#     document.add_paragraph("That's all from us for now. See you next month!", style=text_style)
#     document.add_paragraph("Yours faithfully,\nSocial Impact Team", style=text_style)
#     document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT
#
#
# def generate_bar_chart(dataset, color, name="fig_temp.png"):
#     fig, ax = plt.subplots(figsize=(16, 9))
#     labels = []
#     for x in dataset.index:
#         if len(x) > 30:
#             labels.append(x[:x[:30].rfind(" ")] + "\n" + x[x[:30].rfind(" ") + 1:])
#         else:
#             labels.append(x)
#     print(dataset["values"])
#     bars = ax.barh(width=dataset["values"], y=labels, alpha=0.98, facecolor=COLORS[color])
#     for index, bar in enumerate(bars):
#         ax.text(
#             dataset["values"][index] - 0.1 * dataset.min(),
#             bar.get_y() + bar.get_height() / 2,
#             int(bar.get_width()),
#             ha="right",
#             va="center",
#             color=COLORS["white"],
#         )
#
#     ax.set_axisbelow(True)
#     ax.grid(which="major", axis="x", alpha=0.5)
#
#     plt.tight_layout()
#     plt.savefig(name)
#
#
# def generate_map_chart(dataset, legend_title, date_from, date_to, gdf_countries, name="fig_temp.png"):
#     colors = gdf_countries.merge(dataset[["country", "values", "color"]], how="left",
#                                            left_on="ADMIN", right_on="country")
#     colors["color"] = colors["color"].cat.add_categories(["black", "grey"]).fillna(colors["base_color"])
#
#     fig, ax = plt.subplots(1, 1, figsize=(16, 9))
#     for c in colors["color"].unique():
#         colors[colors["color"] == c].plot(ax=ax, color=COLORS[c], linewidth=0.1, ec="#CCCCCC")
#
#     # Limits, clearing ticks
#     ax.set_aspect(1.)
#     ax.set_xlim(-15, 65)
#     ax.set_ylim(30, 75)
#     ax.set_xticks([])
#     ax.set_yticks([])
#
#     # Delete spines
#     for d in ["left", "top", "right", "bottom"]:
#         ax.spines[d].set_visible(False)
#
#     # CUSTOM LEGEND
#     # Calculating quartile borders
#     borders = dataset["values"].quantile([0, 0.25, 0.5, 0.75, 1], interpolation="nearest").values
#     borders[0] = 0
#     # Creating labels
#     legend_labels = ["no submission"]
#     for i in range(len(borders) - 1):
#         legend_labels.append(f"{int(borders[i] + 1):,} - {int(borders[i + 1]):,}")
#     # Legend
#     custom_lines = [Line2D([0], [0], color=COLORS["darkblue"], lw=4),
#                     Line2D([0], [0], color=COLORS["cyan"], lw=4),
#                     Line2D([0], [0], color=COLORS["orange"], lw=4),
#                     Line2D([0], [0], color=COLORS["magenta"], lw=4),
#                     Line2D([0], [0], color=COLORS["black"], lw=4)
#                     ]
#
#     ax.legend(custom_lines, legend_labels[-1::-1], title=f"{legend_title}\n{date_from} - {date_to}",
#               loc="upper right")
#
#     plt.tight_layout()
#     plt.savefig(name)
#
#
# def generate_doc_report(date_from, date_to, chart_datasets, map_datasets, datasets):
#     gdf_countries = gpd.read_file(COUNTRIES_MAP_PATH)
#     gdf_countries = gdf_countries.merge(datasets.organisations[["country", "base_color"]], how="left",
#                                         left_on="ADMIN", right_on="country").drop("country", axis=1)
#     gdf_countries["base_color"] = gdf_countries["base_color"].fillna("grey")
#     print("\n Generating document report...")
#     # Load titles
#     with open(GRAPH_TITLES_PATH, "r") as f:
#         titles = f.read().splitlines()
#     # Document texts
#     with open(GRAPH_TEXTS_PATH, "r") as f:
#         texts = f.read().replace("\n", "")
#     text_values = {
#         "top_cause": chart_datasets["top_causes"].index[-1], "top_cause_count": chart_datasets["top_causes"]['values'][-1],
#         "top_section": chart_datasets["top_organisers"].index[-1], "top_section_count": chart_datasets["top_organisers"]['values'][-1],
#         "top_events_country": chart_datasets["top_countries"].index[-1],
#         "top_type": chart_datasets["top_categories"].index[-1], "top_type_count": chart_datasets["top_categories"]['values'][-1],
#         "top_goal": chart_datasets["top_sdgs"].index[-1],
#         "top_objective": chart_datasets["top_objectives"].index[-1],
#         "top_country_participants": chart_datasets["top_participants"].index[-1], "top_country_participants_count": chart_datasets["top_participants"]['values'][-1]
#     }
#     graph_texts = texts.format(**text_values).split("$")
#
#     # GRAPH SETTINGS
#     graph_colors = ["darkblue", "magenta", "cyan", "green", "orange"]
#     plt.style.use(GRAPH_STYLE_PATH)
#     # print(plt.rcParams.keys())
#
#     # DOCUMENT SETTINGS
#     # Generate document class
#     document = Document()
#     # Styles
#     styles = change_styles(document.styles)
#     title_style, header_style, text_style = styles
#
#     # DOCUMENT GENERATION
#     # Generating intro
#     generate_doc_intro(document, styles, date_from, date_to, datasets.activities.shape[0], datasets.activities["total_participants"].sum(), datasets.activities["total_participants"].mean())
#     # Generating maps
#     legend_titles = ["Number of submitted events", "Number of submitted participants"]
#     for dataset, title in zip(map_datasets.values(), legend_titles):
#         generate_map_chart(dataset, title, date_from, date_to, gdf_countries)
#         generate_doc_section(document, styles)
#     # Generating graphs
#     for i, (title, dataset, text) in enumerate(zip(titles, list(chart_datasets.values()), graph_texts)):
#         color = graph_colors[i % len(graph_colors)]
#         generate_bar_chart(dataset, color)
#         generate_doc_section(document, styles, text, title)
#     # Generating the final word
#     generate_doc_outro(document, styles)
#     # Saving the document
#     document.save(DOCUMENT_SAVE_PATH.joinpath(f'ESN Activities Report from {date_from} to {date_to}.docx'))
#     # Deleting the temp file
#     Path("fig_temp.png").unlink()
#######################################################################################################################
#######################################################################################################################
#######################################################################################################################


class OrganisationRecord:
    """
    Class representing one organisation
    """
    def __init__(self, record, is_no):
        self.name = record["label"]
        self._cc = record["cc"]
        self.level = "national" if is_no else "local"
        self.country = record["country"]

        self.organisation_record = pd.DataFrame(data={"name": [self.name], "cc": [self._cc], "level": [self.level],
                                                      "country": [self.country]})


class ActivityRecord:
    """
    Class representing one record of activities
    """
    def __init__(self, record):
        self.id = record["id"]
        self.uuid = record["uuid"]
        self.type = record["type"]
        self.created = datetime.fromtimestamp(record["created"])
        self.last_change = datetime.fromtimestamp(record["changed"])
        self.title = record["title"]
        self.no_iso = record["country"]
        self.no_name = "" # need to search the iso enum
        self.start_date = datetime.fromtimestamp(record["dates"]["start"])
        self.end_date = datetime.fromtimestamp(record["dates"]["end"])
        self.description = BSoup(record["description"], "html.parser").get_text()
        self.goal = BSoup(record["goal"], "html.parser").get_text()
        self.main_organiser = record["organisers"][0]
        self.other_organisers = record["organisers"][1:]
        self.url = record["url"]
        self.local_participants = record["participants"]["local"]
        self.intl_participants = record["participants"]["international"]
        self.coordinators = record["participants"]["coordinators"]
        # Physical Event
        if self.type == "physical" and isinstance(record["physical_data"], dict):
            self.location_country_iso = record["physical_data"]["country"]
            self.location_country = "" # need to search the iso enum
            self.location_city = record["physical_data"]["city"] if len(record["physical_data"]["city"]) > 3 else ""
        # Online Event or missing data
        else:
            self.location_country_iso = self.no_iso
            self.location_country = self.no_name
            self.location_city = ""
        # Causes, SDGs, Objectives, Categories
        self.causes = [cause["name"] for cause in record["causes"] if cause["is_main_cause"]]
        self.sdgs = [sdg["slug"] for sdg in record["sdgs"]]
        self.objectives = [obj["name"] for obj in record["objectives"]]
        self.categories = [cat["name"] for cat in record["categories"]]

    # DataFrames
        # General info
        self.general_record = pd.DataFrame(
            data=[[self.id, self.uuid, self.type, self.created, self.last_change, self.title, self.no_iso, self.no_name,
                self.start_date, self.end_date, self.description, self.goal, self.main_organiser, self.other_organisers,
                self.url, self.local_participants, self.intl_participants, self.coordinators, self.location_country_iso,
                self.location_country, self.location_city]],
            columns=["id", "uuid", "type", "created", "last_change", "title", "no_iso", "no_name", "start_date",
                     "end_date", "description", "goal", "main_organiser", "other_organisers", "url",
                     "local_participants", "intl_participants", "coordinators", "location_country_iso",
                     "location_country", "location_city"])
        # Other Organisers
        self.organisers_record = pd.DataFrame(
            data={"id": [self.id] * len(record["organisers"]),
                  "start_date":  [self.start_date] * len(record["organisers"]),
                  "organiser": record["organisers"]}
        )
        # Causes
        self.causes_record = pd.DataFrame(
            data={"id": [self.id] * len(self.causes),
                  "start_date":  [self.start_date] * len(self.causes),
                  "cause": self.causes},
        )
        # SDGs
        self.sdgs_record = pd.DataFrame(
            data={"id": [self.id] * len(self.sdgs),
                  "start_date":  [self.start_date] * len(self.sdgs),
                  "sdg": self.sdgs},
        )
        # Objectives
        self.objectives_record = pd.DataFrame(
            data={"id": [self.id] * len(self.objectives),
                  "start_date":  [self.start_date] * len(self.objectives),
                  "objective": self.objectives},
        )
        # Categories
        self.categories_record = pd.DataFrame(
            data={"id": [self.id] * len(self.categories),
                  "start_date":  [self.start_date] * len(self.categories),
                  "category": self.categories},
        )


class ChartDatasets:
    """
    Class representing chart and map generator
    """
    def __init__(self, records, date_from, date_to, top=5):
        self.activities = records["activities"][(records["activities"]["start_date"] >= date_from) &
                                                (records["activities"]["start_date"] <= date_to)]
        self.activities["total_participants"] = self.activities[["local_participants", "intl_participants", "coordinators"]].sum(axis=1)
        self.activity_type = records["categories"][(records["categories"]["start_date"] >= date_from) &
                                                (records["categories"]["start_date"] <= date_to)]
        self.causes = records["causes"][(records["causes"]["start_date"] >= date_from) &
                                                (records["causes"]["start_date"] <= date_to)]
        self.goals = records["sdgs"][(records["sdgs"]["start_date"] >= date_from) &
                                                (records["sdgs"]["start_date"] <= date_to)]
        self.objectives = records["objectives"][(records["objectives"]["start_date"] >= date_from) &
                                                (records["objectives"]["start_date"] <= date_to)]
        self.organisers = records["organisers"][(records["organisers"]["start_date"] >= date_from) &
                                                (records["organisers"]["start_date"] <= date_to)]
        self.organisations = records["organisations"]
        self.organisations["base_color"] = "black"

        # Analysing loaded data
        self.causes_agg = self.causes.groupby("cause").agg({"cause": "count"}).rename(columns={"cause": "values"}) \
            .sort_values("values")

        self.organisers_agg = self.organisers.groupby(["organiser"]).agg({"organiser": "count"}) \
            .rename(columns={"organiser": "values"}).sort_values("values")

        self.countries_agg = self.activities[["main_organiser"]].merge(
            self.organisations[["name", "country", "national_organisation"]],
            how="left", left_on="main_organiser",
            right_on="name") \
            .groupby(["national_organisation", "country"]).agg({"national_organisation": "count"}).rename(columns={"national_organisation": "values"}) \
            .sort_values("values")
        self.countries_agg["color"] = pd.qcut(self.countries_agg["values"], q=4,
                                              labels=["magenta", "orange", "cyan", "darkblue"])

        self.types_agg = self.activity_type.groupby("category").agg({"category": "count"}) \
            .rename(columns={"category": "values"})

        self.goals_agg = self.goals.groupby("sdg").agg({"sdg": "count"}).rename(columns={"sdg": "values"})
        self.objectives_agg = self.objectives.groupby("objective").agg({"objective": "count"}) \
            .rename(columns={"objective": "values"})

        self.participants_agg = self.activities[["main_organiser", "total_participants"]] \
            .merge(self.organisations[["name", "country", "national_organisation"]], how="left", left_on="main_organiser",
                   right_on="name").groupby(["national_organisation", "country"]) \
            .agg({"total_participants": "sum"}).rename(columns={"total_participants": "values"}).sort_values("values")
        self.participants_agg["color"] = pd.qcut(self.participants_agg["values"], q=4,
                            labels=["magenta", "orange", "cyan", "darkblue"])

        self.top_data = {"top_causes": self.causes_agg.sort_values("values")[-top:],
                        "top_organisers": self.organisers_agg[self.organisers_agg["values"] >= self.organisers_agg["values"][-top]],
                        "top_countries": self.countries_agg[self.countries_agg["values"] >= self.countries_agg["values"][-top]],
                        "top_categories": self.types_agg.sort_values("values")[-top:],
                        "top_sdgs": self.goals_agg.sort_values("values")[-top:],
                        "top_objectives": self.objectives_agg.sort_values("values")[-top:],
                        "top_participants": self.participants_agg[self.participants_agg["values"] >= self.participants_agg["values"][-top]]}

        self.map_data = {"event_count": self.countries_agg.reset_index(),
                      "participant_count": self.participants_agg.reset_index()
        }

        self.text_values = {
            "top_cause": self.top_data["top_causes"].index[-1],
            "top_cause_count": self.top_data["top_causes"]['values'][-1],
            "top_section": self.top_data["top_organisers"].index[-1],
            "top_section_count": self.top_data["top_organisers"]['values'][-1],
            "top_events_country": self.top_data["top_countries"].index[-1][0],
            "top_type": self.top_data["top_categories"].index[-1],
            "top_type_count": self.top_data["top_categories"]['values'][-1],
            "top_goal": self.top_data["top_sdgs"].index[-1],
            "top_objective": self.top_data["top_objectives"].index[-1],
            "top_country_participants": self.top_data["top_participants"].index[-1][0],
            "top_country_participants_count": self.top_data["top_participants"]['values'][-1]
        }

        for name, dataset in self.top_data.items():
            dataset.to_csv(FILE_DIRECTORY.joinpath(f"{name}_{STAMP}.csv"))


class DocReport:
    def __init__(self, records, document, date_from=None, date_to=None):
        self.document = document
        self.document_styles = self.document.styles
        self.title_style = None
        self.heading_style = None
        self.text_style = None
        if date_from is None or date_to is None:
            now = datetime.now()
            if now.day > 15:
                date_to = datetime(now.year, now.month, 1) - relativedelta(seconds=1)
            else:
                date_to = datetime(now.year, now.month, 1) - relativedelta(months=1, seconds=1)
            date_from = date_to - relativedelta(months=1, seconds=-1)

            print(f"Date not set, creating a report from the period {date_from.strftime('%d %b %Y')} "
                  f"- {date_to.strftime('%d %b %Y')}.")
        self.chart_datasets = ChartDatasets(records, date_from, date_to)

        # Title text
        with open(TITLE_TEXT_PATH, "r") as title_text:
            self.title = title_text.read()
        # Intro text
        with open(INTRO_TEXT_PATH, "r") as intro_texts:
            self.intro_texts = intro_texts.read()\
                .format(date_from=date_from.strftime("%d %b %Y"), date_to=date_to.strftime("%d %b %Y"))\
                .replace("\n", "").split("$")
        # Titles for each section
        with open(GRAPH_TITLES_PATH, "r") as title_file:
            self.section_titles = title_file.read().split('\n')
        # Texts for each section
        with open(GRAPH_TEXTS_PATH, "r") as text_file:
            self.texts = text_file.read().format(** self.chart_datasets.text_values).replace("\n", "").split('$')
        # Outro text
        with open(OUTRO_TEXT_PATH, "r") as outro_texts:
            self.outro_texts = outro_texts.read().replace("\n","").split("$")

    def set_styles(self):
        self.title_style = self.document_styles.add_style("TitleStyle", WD_STYLE_TYPE.PARAGRAPH)
        self.title_style.font.name = "Kelson Sans"
        self.title_style.font.size = Pt(24)
        self.title_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        self.title_style.paragraph_format.keep_with_next = True
        self.heading_style = self.document_styles.add_style("HeaderStyle", WD_STYLE_TYPE.PARAGRAPH)
        self.heading_style.font.name = "Lato"
        self.heading_style.font.size = Pt(16)
        self.heading_style.font.bold = True
        self.heading_style.font.color.rgb = RGBColor(0x2E, 0x31, 0x92)
        self.heading_style.paragraph_format.keep_with_next = True
        self.text_style = self.document_styles.add_style("TextStyle", WD_STYLE_TYPE.PARAGRAPH)
        self.text_style.font.name = "Lato"
        self.text_style.font.size = Pt(11)
        self.text_style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        self.text_style.paragraph_format.keep_with_next = True
        self.text_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def generate_report(self):
        if not REPORT_DIRECTORY.exists():
            REPORT_DIRECTORY.mkdir()
        # Add the title
        self.document.add_paragraph(self.title, style=self.title_style)
        # Add the intro text
        paragraph = self.document.add_paragraph(style=self.text_style)
        paragraph.add_run(self.intro_texts[0])
        runner = paragraph.add_run(self.intro_texts[1])
        runner.bold = True
        runner = paragraph.add_run(self.intro_texts[2])
        runner.bold = True
        runner.italic = True
        # Add sections
        for title, text in zip(self.section_titles, self.texts):
            # Create the category heading
            self.document.add_paragraph(title, style=self.heading_style)
            # Add the text to each category
            self.document.add_paragraph(text, style=self.text_style)
        # Add the outro text
        for text in self.outro_texts:
            self.document.add_paragraph(text, style=self.text_style)
        # Save the document
        self.document.save(REPORT_DIRECTORY.joinpath(f"{STAMP}.docx"))


def pipeline():
    print(f"Accessing {ACTIVITIES_URL} to fetch updated information.")
    response = requests.get(url=ACTIVITIES_URL, headers=ACTIVITIES_HEADER)
    try:
        activities_data = response.json()["data"]
    except Exception as e:
        print(f"Exception {e} occured. The status code is {response.status_code}")
    print(f"Request complete. Proceeding with processing individual records...")
    counter = 1
    for _, record in activities_data.items():
        if counter == 1:
            activity_record = ActivityRecord(record)
            activity_record.general_record.to_csv(FILE_DIRECTORY.joinpath(f"activities_{STAMP}.csv"), mode="w", index=False, header=True)
            activity_record.organisers_record.to_csv(FILE_DIRECTORY.joinpath(f"organisers_{STAMP}.csv"), mode="w", index=False, header=True)
            activity_record.causes_record.to_csv(FILE_DIRECTORY.joinpath(f"causes_{STAMP}.csv"), mode="w", index=False, header=True)
            activity_record.sdgs_record.to_csv(FILE_DIRECTORY.joinpath(f"sdgs_{STAMP}.csv"), mode="w", index=False, header=True)
            activity_record.objectives_record.to_csv(FILE_DIRECTORY.joinpath(f"objectives_{STAMP}.csv"), mode="w", index=False, header=True)
            activity_record.categories_record.to_csv(FILE_DIRECTORY.joinpath(f"categories_{STAMP}.csv"), mode="w", index=False, header=True)
        else:
            activity_record = ActivityRecord(record)
            activity_record.general_record.to_csv(FILE_DIRECTORY.joinpath(f"activities_{STAMP}.csv"), mode="a", index=False, header=False)
            activity_record.organisers_record.to_csv(FILE_DIRECTORY.joinpath(f"organisers_{STAMP}.csv"), mode="a", index=False, header=False)
            activity_record.causes_record.to_csv(FILE_DIRECTORY.joinpath(f"causes_{STAMP}.csv"), mode="a", index=False, header=False)
            activity_record.sdgs_record.to_csv(FILE_DIRECTORY.joinpath(f"sdgs_{STAMP}.csv"), mode="a", index=False, header=False)
            activity_record.objectives_record.to_csv(FILE_DIRECTORY.joinpath(f"objectives_{STAMP}.csv"), mode="a", index=False, header=False)
            activity_record.categories_record.to_csv(FILE_DIRECTORY.joinpath(f"categories_{STAMP}.csv"), mode="a", index=False, header=False)

        if (counter % 500) == 0:
            print(f" Records processed: {counter}...")
        counter += 1

    print(f"Processing done! Total of {counter} activities saved.")
    print("Proceeding with downloading the organisations information...")

    # activities = pd.read_csv(FILE_DIRECTORY.joinpath(f"activities_{STAMP}.csv"))
    # organisers = pd.read_csv(FILE_DIRECTORY.joinpath(f"organisers_{STAMP}.csv"))
    # causes = pd.read_csv(FILE_DIRECTORY.joinpath(f"causes_{STAMP}.csv"))
    # sdgs = pd.read_csv(FILE_DIRECTORY.joinpath(f"sdgs_{STAMP}.csv"))
    # objectives = pd.read_csv(FILE_DIRECTORY.joinpath(f"objectives_{STAMP}.csv"))
    # categories = pd.read_csv(FILE_DIRECTORY.joinpath(f"categories_{STAMP}.csv"))


def get_organisations_info():
    counter = 1
    for url, is_no in zip([SECTIONS_URL, COUNTRIES_URL], [False, True]):
        r = requests.get(url)
        records = r.json()
        for record in records:
            if counter == 1:
                organisation_record = OrganisationRecord(record, is_no).organisation_record
            else:
                organisation_record = pd.concat([
                    organisation_record,
                    OrganisationRecord(record, is_no).organisation_record
                ])
            counter += 1

    organisations_nos = organisation_record[organisation_record.level == "national"][["name", "cc"]] \
        .rename(columns={"name": "national_organisation"})
    organisations = organisation_record.merge(organisations_nos, on=["cc", "cc"], how="left").drop(columns=["cc"])
    organisations.to_csv(FILE_DIRECTORY.joinpath(f"organisations_{STAMP}.csv"), index=False)

def load_datasets():
    names = ["activities", "organisers", "causes", "sdgs", "objectives", "categories", "organisations"]
    data = {}
    for name in names:
        data[name] = pd.read_csv(FILE_DIRECTORY.joinpath(f"{name}_{STAMP}.csv"))
        if "start_date" in data[name].columns:
            data[name]["start_date"] = pd.to_datetime(data[name]["start_date"])

    return data


FILE_DIRECTORY = Path("files")
FILE_DIRECTORY.mkdir(exist_ok=True, parents=True)
COUNTRIES_MAP_PATH = Path("report_files", "countries.geojson")
GRAPH_TITLES_PATH = Path("report_files", "titles_graphs.txt")
GRAPH_TEXTS_PATH = Path("report_files", "texts_graphs.txt")
GRAPH_STYLE_PATH = Path("report_files", "style_graphs.mplstyle")
TITLE_TEXT_PATH = Path("report_files", "title_text.txt")
INTRO_TEXT_PATH = Path("report_files", "intro_texts.txt")
OUTRO_TEXT_PATH = Path("report_files", "outro_texts.txt")
REPORT_DIRECTORY = Path("reports")
STAMP = "test"
ACTIVITIES_HEADER = {"X-Api-Key": "014709cfb534266769769522ac9ebaab"}
ACTIVITIES_URL = "https://api.erasmusgeneration.org/api/v1/static/activities"
SECTIONS_URL = "https://accounts.esn.org/api/v1/sections"
COUNTRIES_URL = "https://accounts.esn.org/api/v1/countries"
# ESN COLORS
COLORS = {
    "darkblue": "#2E3192",
    "cyan": "#00AEEF",
    "magenta": "#EC008C",
    "green": "#7AC143",
    "orange": "#F47B20",
    "black": "#000000",
    "white": "#FFFFFF",
    "grey": "#888888"
}


if __name__ == '__main__':
    # pipeline()
    get_organisations_info()
    data = load_datasets()
    document = Document()
    report = DocReport(data, document)
    report.set_styles()
    report.generate_report()
